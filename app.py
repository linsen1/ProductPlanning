"""
AI 策划智能体 — 一期：产品策划方案（两段式）
技术栈：Streamlit + LangGraph + Agno + Chroma RAG + 千问3 + 博查搜索
流程：填表 → 博查搜竞品 → AI给USP方向 → 策划人选方向 → AI生成完整方案
"""
from __future__ import annotations
import hashlib
import io
import math
import os
import json
import re
import requests
import tempfile
import uuid
from pathlib import Path
from typing import Dict, List, TypedDict, Optional

# Some cloud images can resolve an old opentelemetry-proto package with a
# newer protobuf runtime before Chroma imports. Use the pure-Python fallback so
# the app can still start while requirements keep the dependency stack aligned.
os.environ.setdefault("PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION", "python")

import streamlit as st
import streamlit.components.v1 as components
from pathlib import Path as _Path
from dotenv import load_dotenv
from langchain_core.embeddings import Embeddings
from langgraph.graph import StateGraph, END
from langchain_core.documents import Document
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from agno.agent import Agent
from agno.models.dashscope import DashScope

import config
from prompts import (
    STRATEGY_PROMPT,
    PRODUCT_PLAN_PROMPT,
    PRODUCT_SYSTEM_PROMPT,
    build_retrieval_query,
    format_retrieved_cases,
    format_kb_blocks,
    format_bocha_results,
)

load_dotenv(dotenv_path=_Path(__file__).parent / ".env", override=True)


# ── LLM / Embedding 工厂 ──────────────────────────────────────────────────────
def get_llm(temperature: float = 0.3) -> ChatOpenAI:
    if config.LLM_BACKEND == "openai":
        return ChatOpenAI(model=config.OPENAI_MODEL, api_key=config.OPENAI_API_KEY, temperature=temperature)
    return ChatOpenAI(
        model=config.QWEN_MODEL, api_key=config.DASHSCOPE_API_KEY,
        base_url=config.QWEN_BASE_URL, temperature=temperature,
        extra_body={"enable_thinking": False},
    )

def get_agno_model(temperature: float = 0.5) -> DashScope:
    return DashScope(
        id=config.QWEN_MODEL, api_key=config.DASHSCOPE_API_KEY,
        base_url=config.QWEN_BASE_URL, temperature=temperature,
        enable_thinking=False,
    )

_embedding_fallback_warning_shown = False


class LocalHashEmbeddings(Embeddings):
    """Deterministic local fallback used when remote embedding access is denied."""

    def __init__(self, dimensions: int = 1024):
        self.dimensions = dimensions

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [self._embed(text) for text in texts]

    def embed_query(self, text: str) -> List[float]:
        return self._embed(text)

    def _embed(self, text: str) -> List[float]:
        vec = [0.0] * self.dimensions
        tokens = self._tokens(text or "")
        if not tokens:
            return vec
        for token, weight in tokens:
            digest = hashlib.blake2b(token.encode("utf-8"), digest_size=8).digest()
            idx = int.from_bytes(digest[:4], "big") % self.dimensions
            sign = 1.0 if digest[4] & 1 else -1.0
            vec[idx] += sign * weight
        norm = math.sqrt(sum(v * v for v in vec))
        if not norm:
            return vec
        return [v / norm for v in vec]

    @staticmethod
    def _tokens(text: str) -> List[tuple[str, float]]:
        lowered = text.lower()
        tokens: List[tuple[str, float]] = []
        tokens.extend((m.group(0), 1.0) for m in re.finditer(r"[a-z0-9]+|[\u4e00-\u9fff]", lowered))
        compact = re.sub(r"\s+", "", lowered)
        chars = [c for c in compact if c.strip()]
        for n, weight in ((2, 1.5), (3, 1.2)):
            if len(chars) >= n:
                tokens.extend(("".join(chars[i:i + n]), weight) for i in range(len(chars) - n + 1))
        return tokens


class ResilientEmbeddings(Embeddings):
    def __init__(self, remote: Embeddings, fallback: Embeddings, model_name: str):
        self.remote = remote
        self.fallback = fallback
        self.model_name = model_name
        self.remote_disabled = False

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        if self.remote_disabled:
            return self.fallback.embed_documents(texts)
        try:
            return self.remote.embed_documents(texts)
        except Exception as exc:
            self.remote_disabled = True
            _warn_embedding_fallback(exc, self.model_name)
            return self.fallback.embed_documents(texts)

    def embed_query(self, text: str) -> List[float]:
        if self.remote_disabled:
            return self.fallback.embed_query(text)
        try:
            return self.remote.embed_query(text)
        except Exception as exc:
            self.remote_disabled = True
            _warn_embedding_fallback(exc, self.model_name)
            return self.fallback.embed_query(text)


def _warn_embedding_fallback(exc: Exception, model_name: str) -> None:
    global _embedding_fallback_warning_shown
    if _embedding_fallback_warning_shown:
        return
    _embedding_fallback_warning_shown = True
    msg = str(exc)
    if len(msg) > 220:
        msg = msg[:220] + "..."
    st.warning(
        f"DashScope embedding 模型 `{model_name}` 暂不可用，已自动切换到本地知识库向量。"
        f"原始错误：{msg}"
    )


def get_embeddings() -> Embeddings:
    fallback = LocalHashEmbeddings()
    if config.EMBEDDING_BACKEND.lower() == "local" or not config.DASHSCOPE_API_KEY:
        return fallback
    remote = OpenAIEmbeddings(
        model=config.EMBEDDING_MODEL, api_key=config.DASHSCOPE_API_KEY,
        base_url=config.QWEN_BASE_URL, check_embedding_ctx_length=False,
        chunk_size=8,   # text-embedding-v4 单批上限10，留余量设8
    )
    return ResilientEmbeddings(remote, fallback, config.EMBEDDING_MODEL)


# ── 博查搜索 ──────────────────────────────────────────────────────────────────
def bocha_search(query: str, count: int = None, freshness: str = None) -> list[dict]:
    """调用博查 Web Search API，返回搜索结果列表"""
    if not config.bocha_enabled():
        return []
    try:
        resp = requests.post(
            config.BOCHA_API_URL,
            headers={
                "Authorization": f"Bearer {config.BOCHA_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "query": query,
                "freshness": freshness or config.BOCHA_FRESHNESS,
                "summary": True,
                "count": count or config.BOCHA_SEARCH_COUNT,
            },
            timeout=20,
        )
        resp.raise_for_status()
        data = resp.json()
        return data.get("data", {}).get("webPages", {}).get("value", []) \
            or data.get("webPages", {}).get("value", [])
    except Exception as e:
        st.warning(f"⚠️ 博查搜索异常（{e}），将跳过竞品搜索继续生成")
        return []


def bocha_ai_search(query: str, count: int = None) -> tuple[str, list[dict]]:
    """调用博查 AI Search，返回 (综合长答案, 来源网页列表)。失败则返回 ('', [])"""
    if not config.bocha_enabled():
        return "", []
    try:
        resp = requests.post(
            config.BOCHA_AISEARCH_URL,
            headers={
                "Authorization": f"Bearer {config.BOCHA_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "query": query,
                "freshness": config.BOCHA_FRESHNESS,
                "answer": True,
                "stream": False,
                "count": count or config.BOCHA_SEARCH_COUNT,
            },
            timeout=40,
        )
        resp.raise_for_status()
        d = resp.json()
        msgs = (d.get("data", {}) or {}).get("messages") or d.get("messages") or []
        answer, pages = "", []
        for m in msgs:
            mtype = m.get("type", "")
            content = m.get("content", "")
            if mtype == "answer":
                answer += content if isinstance(content, str) else json.dumps(content, ensure_ascii=False)
            elif mtype == "source":
                obj = content
                if isinstance(obj, str):
                    try:
                        obj = json.loads(obj)
                    except Exception:
                        obj = {}
                vals = (obj.get("value")
                        or (obj.get("webPages", {}) or {}).get("value")
                        or [])
                pages += [v for v in vals if isinstance(v, dict)]
        return answer.strip(), pages
    except Exception as e:
        st.warning(f"⚠️ 博查 AI 搜索异常（{e}），改用普通搜索")
        return "", []


# ── 向量库管理 ────────────────────────────────────────────────────────────────
def _is_streamlit_cloud() -> bool:
    cwd = Path.cwd().resolve()
    return str(cwd).startswith("/mount/src") or bool(os.getenv("STREAMLIT_SHARING_MODE"))


def _vectorstore_persistent_enabled() -> bool:
    mode = config.VECTOR_DB_MODE.lower()
    if mode in ("memory", "in-memory", "ephemeral"):
        return False
    if mode in ("persistent", "persist", "disk"):
        return True
    return not _is_streamlit_cloud()


def _vector_db_path() -> Path:
    path = Path(config.DB_DIR)
    if _is_streamlit_cloud() and not path.is_absolute():
        return Path(tempfile.gettempdir()) / "projectai" / path.name
    return path


def _memory_collection_name() -> str:
    return f"projectai-kb-{uuid.uuid4().hex[:16]}"


def _kb_manifest_path() -> Path:
    return _vector_db_path() / "kb_manifest.json"

def _kb_manifest() -> list[dict]:
    kb_root = Path(config.KB_DIR)
    if not kb_root.exists():
        return []
    manifest = []
    for p in sorted(kb_root.glob("**/*.md")):
        stat = p.stat()
        manifest.append({
            "path": str(p.relative_to(kb_root)),
            "mtime_ns": stat.st_mtime_ns,
            "size": stat.st_size,
        })
    return manifest

def _load_kb_manifest() -> list[dict]:
    path = _kb_manifest_path()
    if not path.exists():
        return []
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return []

def _save_kb_manifest(manifest: list[dict]) -> None:
    path = _kb_manifest_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

def _kb_manifest_changed() -> bool:
    return _load_kb_manifest() != _kb_manifest()

def _kb_manifest_key() -> str:
    return json.dumps(_kb_manifest(), ensure_ascii=False, sort_keys=True)

# ── front-matter 解析（P2：供检索硬过滤）─────────────────────────────────────
def _parse_frontmatter(text: str) -> tuple[dict, str]:
    """解析顶部 YAML front-matter（轻量实现，不依赖 yaml 库），返回 (meta, body)"""
    meta: dict = {}
    t = text.lstrip()
    if t.startswith("---"):
        end = t.find("\n---", 3)
        if end != -1:
            fm = t[3:end].strip()
            body = t[end + 4:].lstrip("\n")
            for line in fm.splitlines():
                if ":" in line:
                    k, v = line.split(":", 1)
                    meta[k.strip()] = v.strip()
            return meta, body
    return meta, text

def _norm_meta(meta: dict, source: str) -> dict:
    """归一化为 Chroma 可存的标量 metadata"""
    return {
        "source":       source,
        "project_type": meta.get("project_type", "") or "",
        "cert":         meta.get("cert", "") or "",
        "layer":        meta.get("layer", "") or "",
        "priority":     meta.get("priority", "") or "",
        "applies_to":   meta.get("applies_to", "") or "",
    }

def build_vectorstore():
    kb_root = Path(config.KB_DIR)
    docs = []
    for p in sorted(kb_root.glob("**/*.md")):
        if p.name.startswith("README"):   # 文档说明类不入 RAG
            continue
        text = p.read_text(encoding="utf-8")
        meta, _body = _parse_frontmatter(text)
        docs.append(Document(
            page_content=text,
            metadata=_norm_meta(meta, str(p.relative_to(kb_root))),
        ))
    if not docs:
        return None, 0, 0
    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=100)
    splits = splitter.split_documents(docs)
    persist = _vectorstore_persistent_enabled()
    try:
        if persist:
            _vector_db_path().mkdir(parents=True, exist_ok=True)
            vs = Chroma.from_documents(
                splits,
                embedding=get_embeddings(),
                persist_directory=str(_vector_db_path()),
            )
            _save_kb_manifest(_kb_manifest())
        else:
            vs = Chroma.from_documents(
                splits,
                embedding=get_embeddings(),
                collection_name=_memory_collection_name(),
            )
    except Exception as exc:
        if not persist:
            raise
        reset_vectorstore()
        st.warning(f"持久化向量库初始化失败，已自动切换到内存向量库。原始错误：{str(exc)[:220]}")
        vs = Chroma.from_documents(
            splits,
            embedding=get_embeddings(),
            collection_name=_memory_collection_name(),
        )
    return vs, len(docs), len(splits)

def load_vectorstore():
    return Chroma(persist_directory=str(_vector_db_path()), embedding_function=get_embeddings())

def reset_vectorstore():
    import shutil
    db_path = _vector_db_path()
    if db_path.exists():
        shutil.rmtree(db_path, ignore_errors=True)

def _persisted_count() -> int:
    import sqlite3
    db = _vector_db_path() / "chroma.sqlite3"
    if not db.exists(): return 0
    try:
        con = sqlite3.connect(f"file:{db}?mode=ro", uri=True)
        tables = {r[0] for r in con.execute("SELECT name FROM sqlite_master WHERE type='table'")}
        count = con.execute("SELECT COUNT(*) FROM embeddings").fetchone()[0] if "embeddings" in tables else 0
        con.close()
        return count
    except Exception:
        return 0

@st.cache_resource(show_spinner="📚 加载产品案例知识库...")
def ensure_vectorstore(manifest_key: str):
    if not _vectorstore_persistent_enabled():
        return build_vectorstore()
    if (
        not _vector_db_path().exists()
        or not os.listdir(_vector_db_path())
        or _persisted_count() == 0
        or _kb_manifest_changed()
    ):
        reset_vectorstore()
        return build_vectorstore()
    vs = load_vectorstore()
    try:
        if vs._collection.count() == 0:
            del vs; reset_vectorstore()
            return build_vectorstore()
    except Exception:
        pass
    return vs, None, None

def retrieve_cases(query: str) -> List[str]:
    result = ensure_vectorstore(_kb_manifest_key())
    vs = result[0] if isinstance(result, tuple) else result
    if vs is None: return []
    try:
        return [d.page_content for d in vs.similarity_search(query, k=config.TOP_K)]
    except Exception:
        return []


# ── P2：硬过滤 + 分层软召回 ───────────────────────────────────────────────────
def _iter_kb_files():
    kb_root = Path(config.KB_DIR)
    for p in sorted(kb_root.glob("**/*.md")):
        if p.name.startswith("README"):
            continue
        text = p.read_text(encoding="utf-8")
        meta, body = _parse_frontmatter(text)
        yield p, meta, body

def _applies(meta: dict, pt_key: str) -> bool:
    """判断该 KB 文件是否适用于当前项目类型键"""
    if not pt_key or pt_key == "通用":
        return True
    mpt = meta.get("project_type", "") or ""
    app = meta.get("applies_to", "") or ""
    if mpt == pt_key or mpt == "通用":
        return True
    if pt_key in app or "全部" in app:
        return True
    if mpt and (mpt in pt_key or pt_key in mpt):  # 前缀宽松匹配
        return True
    return False

def hard_context(project_type_key: str, certs: List[str]) -> dict:
    """硬命中（不靠相似度，直接读盘）：结构模板 + 证书事实卡 + 合规库"""
    structure, cert_facts, compliance = [], [], []
    certs = certs or []
    for p, meta, body in _iter_kb_files():
        layer = meta.get("layer", "")
        if layer == "structure":
            if meta.get("project_type") == "通用" or _applies(meta, project_type_key):
                structure.append((p.name, body))
        elif layer == "cert_facts":
            c = meta.get("cert", "")
            if (certs and c in certs) or (not certs and _applies(meta, project_type_key)):
                cert_facts.append((c or p.stem, body))
        elif layer == "compliance":
            compliance.append((p.name, body))
    return {"structure": structure, "cert_facts": cert_facts, "compliance": compliance}

def competitor_context(project_type_key: str) -> str:
    """硬命中竞品库（layer==competitors），策略阶段注入，搜索为空时兜底"""
    blocks = []
    for p, meta, body in _iter_kb_files():
        if meta.get("layer", "") == "competitors" and _applies(meta, project_type_key):
            blocks.append((p.stem, body))
    return format_kb_blocks(blocks)

def soft_context(query: str, project_type_key: str) -> dict:
    """分层软召回（embedding + layer 过滤 + project_type 后过滤 + 配额）"""
    result = ensure_vectorstore(_kb_manifest_key())
    vs = result[0] if isinstance(result, tuple) else result
    out = {layer: [] for layer in config.SOFT_RECALL_QUOTAS}
    if vs is None:
        return out
    for layer, quota in config.SOFT_RECALL_QUOTAS.items():
        try:
            docs = vs.similarity_search(query, k=quota * 3, filter={"layer": layer})
        except Exception:
            docs = []
        picked = []
        for d in docs:
            if _applies(d.metadata, project_type_key) and d.page_content not in picked:
                picked.append(d.page_content)
            if len(picked) >= quota:
                break
        if len(picked) < quota:  # 不足则放宽过滤补齐
            for d in docs:
                if d.page_content not in picked:
                    picked.append(d.page_content)
                if len(picked) >= quota:
                    break
        out[layer] = picked
    return out


# ── LangGraph 状态 ────────────────────────────────────────────────────────────
class ProductPlanState(TypedDict, total=False):
    # 基础产品信息
    product_name:        str
    project_type:        str
    plan_mode:           str        # 单产品方案 / 产品体系方案
    certs:               List[str]  # 产品体系模式下选中的证书（驱动多张事实卡 + 分证书简章）
    target_audience:     str
    learning_motivation: str
    core_advantage:      str
    advantage_dims:      str
    price_tier:          str
    extra_info:          str
    # 竞品 & 我方现状（可选）
    competitor_name:     str
    our_pass_rate:       str
    sales_objections:    str
    product_weaknesses:  str
    success_cases:       str
    # 产品结构补充（可选，影响简章/师资/课程说明三章质量）
    reference_price:     str   # 具体售价，如"标准班1680元，协议班2980元"
    teacher_info:        str   # 师资简介，如"王某某，10年教学经验"
    course_structure:    str   # 课程体系，如"基础50课时→强化80课时→冲刺40课时→押题20课时"
    # 中间状态
    retrieval_query:     str
    bocha_results:       str   # 格式化后的搜索文本
    bocha_count:         int   # 博查实际返回去重后条数（诊断用）
    retrieved_cases:     List[str]
    # 分层知识库上下文（P2 检索升级后注入）
    ctx_structure:       str
    ctx_cert_facts:      str
    ctx_skill:           str
    ctx_career:          str
    ctx_copy:            str
    ctx_compliance:      str
    # 第一段输出
    strategy_analysis:   str   # 竞品分析 + USP方向
    # 用户选择
    selected_usp:        str   # 策划人选定的USP方向内容
    # 第二段输出
    product_plan:        str


# ── 第一段 LangGraph：搜索 + 策略分析 ────────────────────────────────────────
def build_query_node(state: ProductPlanState) -> ProductPlanState:
    q = build_retrieval_query(
        state.get("product_name", ""),
        state.get("project_type", ""),
        state.get("target_audience", ""),
    )
    return {"retrieval_query": q}

def bocha_search_node(state: ProductPlanState) -> ProductPlanState:
    """P4 竞品挖深：分证书/赛道多轮搜（定价+卖点+班型+服务+差评），去重后合并"""
    type_kw = state.get("project_type", "").split("—")[-1].split("（")[0].strip() \
        or config.project_type_key(state.get("project_type", ""))
    competitor = state.get("competitor_name", "").strip()
    certs = state.get("certs") or []
    targets = certs[:3] if certs else [type_kw]

    C = config.BOCHA_SEARCH_COUNT
    all_results: list[dict] = []
    ai_answer = ""

    # ⓪ 竞品综合概述 → 走 AI 搜索（长答案 + 来源），深度更高
    if config.BOCHA_USE_AISEARCH:
        focus = competitor or "、".join(targets)
        ai_answer, ai_pages = bocha_ai_search(
            f"{focus} {type_kw} 培训机构 竞品对比：主流机构、卖点、价格档、班型、服务、口碑与常见弱点"
        )
        all_results += ai_pages

    # ① 指定竞品 → 定价/班型/卖点专搜（含口碑）
    if competitor:
        all_results += bocha_search(f"{competitor} {type_kw} 课程 班型 价格 卖点 服务", count=C)
        all_results += bocha_search(f"{competitor} 怎么样 靠谱吗 评价 投诉 退费", count=C, freshness="noLimit")

    # ② 分目标多轮：定价 / 机构对比 / 差评口碑 / 班型服务
    for t in targets:
        all_results += bocha_search(f"{t} 培训 课程 价格 班型 2025 2026", count=C)
        all_results += bocha_search(f"{t} 培训机构 排名 对比 哪个好", count=C, freshness="noLimit")
        all_results += bocha_search(f"{t} 培训机构 评价 差评 退费 投诉 口碑 避坑", count=C, freshness="noLimit")

    # ③ 赛道大盘 — 头部机构主流打法 + 收费
    all_results += bocha_search(f"{type_kw} 头部培训机构 卖点 收费 套路", count=C, freshness="noLimit")

    # 去重（按 url），有效结果计数
    seen, dedup = set(), []
    for r in all_results:
        u = r.get("url", "")
        if u and u in seen:
            continue
        if u:
            seen.add(u)
        dedup.append(r)

    formatted = format_bocha_results(dedup)
    if ai_answer:
        formatted = (
            "## 🤖 AI 综合竞品概述（博查 AI 搜索）\n"
            + ai_answer + "\n\n---\n\n## 🔎 检索原始结果\n" + formatted
        )
    return {"bocha_results": formatted, "bocha_count": len(dedup)}

def gen_strategy_node(state: ProductPlanState) -> ProductPlanState:
    """生成竞品分析 + 2-3个USP方向"""
    pt_key = config.project_type_key(state.get("project_type", ""))
    comp_kb = competitor_context(pt_key)
    prompt = STRATEGY_PROMPT.format(
        product_name        = state.get("product_name", ""),
        project_type        = state.get("project_type", ""),
        target_audience     = state.get("target_audience", ""),
        learning_motivation = state.get("learning_motivation", ""),
        core_advantage      = state.get("core_advantage", ""),
        price_tier          = state.get("price_tier", "待定"),
        our_pass_rate       = state.get("our_pass_rate", "暂无数据"),
        sales_objections    = state.get("sales_objections", "暂无"),
        product_weaknesses  = state.get("product_weaknesses", "暂无"),
        success_cases       = state.get("success_cases", "暂无"),
        bocha_results       = state.get("bocha_results", "（未获取到搜索数据）"),
        competitor_kb       = comp_kb or "（无内置竞品库，请基于行业经验并标注'待核实'）",
    )
    agent = Agent(
        name="strategy-agent",
        model=get_agno_model(0.4),
        description="优路教育资深产品策划顾问，擅长竞品分析和定位策略。",
        instructions=[
            "严格按照模板格式输出，不要添加模板以外的内容。",
            "3个USP方向要有明显差异，不要雷同。",
            "竞品空白分析要具体，不要泛泛而谈。",
            "风险提示要实际，帮策划人规避踩坑。",
        ],
        markdown=True,
    )
    result = agent.run(prompt)
    return {"strategy_analysis": getattr(result, "content", str(result))}

def retrieve_kb_node(state: ProductPlanState) -> ProductPlanState:
    """P2：硬命中（结构/证书卡/合规）+ 分层软召回（技能/变现/金句/案例）"""
    pt_key = config.project_type_key(state.get("project_type", ""))
    query  = state.get("retrieval_query") or build_retrieval_query(
        state.get("product_name", ""),
        state.get("project_type", ""),
        state.get("target_audience", ""),
    )
    certs = state.get("certs") or []
    hard = hard_context(pt_key, certs)
    soft = soft_context(query, pt_key)
    return {
        "ctx_structure":   format_kb_blocks(hard["structure"]),
        "ctx_cert_facts":  format_kb_blocks(hard["cert_facts"]),
        "ctx_compliance":  format_kb_blocks(hard["compliance"]),
        "ctx_skill":       format_retrieved_cases(soft["skill_modules"]),
        "ctx_career":      format_retrieved_cases(soft["career_paths"]),
        "ctx_copy":        format_retrieved_cases(soft["copy_bank"]),
        "retrieved_cases": soft["case"],
    }

def gen_product_plan_node(state: ProductPlanState) -> ProductPlanState:
    """第二段：基于选定USP方向生成完整产品策划方案（单产品 / 产品体系两套模板）"""
    cases_text = format_retrieved_cases(state.get("retrieved_cases", []))
    mode       = state.get("plan_mode", "单产品方案")
    certs      = state.get("certs") or []
    is_system  = "体系" in mode
    template   = PRODUCT_SYSTEM_PROMPT if is_system else PRODUCT_PLAN_PROMPT

    prompt = template.format(
        selected_usp        = state.get("selected_usp", "无特定方向，综合最优"),
        product_name        = state.get("product_name", ""),
        project_type        = state.get("project_type", ""),
        certs_text          = "、".join(certs) if certs else "（未指定证书，按单产品处理）",
        target_audience     = state.get("target_audience", ""),
        learning_motivation = state.get("learning_motivation", ""),
        core_advantage      = state.get("core_advantage", ""),
        advantage_dims      = state.get("advantage_dims", ""),
        price_tier          = state.get("price_tier", "待定"),
        vs_competitor       = state.get("competitor_name", "暂无") or "暂无明确竞品",
        extra_info          = state.get("extra_info", "无"),
        reference_price     = state.get("reference_price", "") or "待策划人确认",
        teacher_info        = state.get("teacher_info", "") or "待策划人提供师资信息",
        course_structure    = state.get("course_structure", "") or "待策划人确认课程体系",
        ctx_structure       = state.get("ctx_structure", "") or "（无结构模板命中）",
        ctx_cert_facts      = state.get("ctx_cert_facts", "") or "（无证书事实卡命中，相关事实请标注待确认）",
        ctx_skill           = state.get("ctx_skill", "") or "（无技能课素材命中，按赛道通用技能模块兜底并标注）",
        ctx_career          = state.get("ctx_career", "") or "（无就业变现素材命中）",
        ctx_copy            = state.get("ctx_copy", "") or "（无金句话术命中）",
        ctx_compliance      = state.get("ctx_compliance", "") or "（无合规库命中，仍须遵守通用合规边界）",
        competitor_analysis = state.get("strategy_analysis", "") or "（无竞品分析数据，竞品矩阵中未掌握项写「暂无公开信息」）",
        retrieved_cases     = cases_text,
    )
    instructions = [
        "严格围绕策划人选定的USP方向展开，不要偏离核心主张。",
        "所有章节都必须完整、达到字数/要素下限，不可省略、不可一句话带过。",
        "证书事实、价格、政策只能用『证书事实卡』内容；卡中没有的写『（参考，待策划人确认）』，禁止编造。",
        "实操技能课必须逐模块展开（来源于技能课素材），这是方案厚度的核心，不能略写。",
        "Slogan、产品优势金句优先从金句话术库选用，再结合产品实际微调。",
        "严格遵守合规库红线，所有高风险表达转为可替代表达。",
        "参考案例只学写法，不照抄。",
    ]
    if is_system:
        instructions.insert(1, f"这是产品体系方案，涉及多个证书（{('、'.join(certs)) or '见输入'}），招生简章需按证书分别出表，每个证书都要覆盖技能课与就业变现。")

    agent = Agent(
        name="product-plan-agent",
        model=get_agno_model(0.5),
        description="优路教育资深产品策划专家，擅长教育行业课程包装和产品定位。",
        instructions=instructions,
        markdown=True,
    )
    result = agent.run(prompt)
    return {"product_plan": getattr(result, "content", str(result))}


# ── 两段式 LangGraph ──────────────────────────────────────────────────────────
@st.cache_resource
def build_strategy_graph():
    wf = StateGraph(ProductPlanState)
    wf.add_node("build_query",    build_query_node)
    wf.add_node("bocha_search",   bocha_search_node)
    wf.add_node("gen_strategy",   gen_strategy_node)
    wf.set_entry_point("build_query")
    wf.add_edge("build_query",  "bocha_search")
    wf.add_edge("bocha_search", "gen_strategy")
    wf.add_edge("gen_strategy", END)
    return wf.compile()

@st.cache_resource
def build_plan_graph():
    wf = StateGraph(ProductPlanState)
    wf.add_node("retrieve_kb",   retrieve_kb_node)
    wf.add_node("generate_plan", gen_product_plan_node)
    wf.set_entry_point("retrieve_kb")
    wf.add_edge("retrieve_kb",   "generate_plan")
    wf.add_edge("generate_plan", END)
    return wf.compile()

def run_strategy(inputs: dict) -> dict:
    return build_strategy_graph().invoke(inputs)

def run_plan(inputs: dict) -> dict:
    return build_plan_graph().invoke(inputs)


# ── Word 导出（P5 重写：markdown → 真 Word，修复表格散架）────────────────────
BRAND_RED_HEX = "C8102E"   # 优路品牌红（占位，可按 VI 精确值调整）

def _md_inline_runs(paragraph, text: str):
    """把含 **加粗** / *斜体* 的行解析为多个 run"""
    import re
    for part in re.split(r"(\*\*.+?\*\*|\*[^*\n]+?\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)

def _is_table_sep(cells: list) -> bool:
    import re
    nonempty = [c for c in cells if c]
    return bool(nonempty) and all(re.fullmatch(r":?-{2,}:?", c) for c in nonempty)

def _split_row(line: str) -> list:
    cells = [c.strip() for c in line.strip().split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells

def export_to_word(product_name: str, content: str) -> Optional[bytes]:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import datetime as _dt

        BRAND = RGBColor.from_string(BRAND_RED_HEX)
        GRAY  = RGBColor(0x80, 0x80, 0x80)

        def _shade(cell, hexcolor):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hexcolor)
            tcPr.append(shd)

        def _fill_cell(cell, text, header=False):
            cell.text = ""
            p = cell.paragraphs[0]
            for k, seg in enumerate(text.split("<br>")):
                if k > 0:
                    p.add_run().add_break()
                _md_inline_runs(p, seg.strip())
            if header:
                _shade(cell, BRAND_RED_HEX)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        doc = DocxDocument()

        # 中文字体
        normal = doc.styles["Normal"]
        normal.font.name = "微软雅黑"
        normal.font.size = Pt(10.5)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 页眉
        try:
            hp = doc.sections[0].header.paragraphs[0]
            hp.text = "优路教育 · 产品策划"
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in hp.runs:
                r.font.size = Pt(8); r.font.color.rgb = GRAY
        except Exception:
            pass

        # 封面
        is_system = "产品体系方案" in content
        subtitle  = "产品体系方案" if is_system else "产品策划方案"
        for _ in range(6):
            doc.add_paragraph("")
        ct = doc.add_paragraph(); ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = ct.add_run(product_name); cr.bold = True
        cr.font.size = Pt(26); cr.font.color.rgb = BRAND
        cs = doc.add_paragraph(); cs.alignment = WD_ALIGN_PARAGRAPH.CENTER
        csr = cs.add_run(subtitle); csr.font.size = Pt(15)
        cd = doc.add_paragraph(); cd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cdr = cd.add_run("优路教育　·　" + _dt.date.today().strftime("%Y年%m月"))
        cdr.font.size = Pt(10.5); cdr.font.color.rgb = GRAY
        doc.add_page_break()

        # 正文逐行/逐块解析
        lines = content.split("\n")
        i, n = 0, len(lines)
        first_h1_skipped = False
        while i < n:
            raw = lines[i].rstrip()
            s = raw.strip()

            # 表格块：连续以 | 开头的行
            if s.startswith("|"):
                block = []
                while i < n and lines[i].strip().startswith("|"):
                    block.append(lines[i]); i += 1
                rows = [c for c in (_split_row(b) for b in block) if not _is_table_sep(c)]
                rows = [r for r in rows if r]
                if rows:
                    ncols = max(len(r) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=ncols)
                    table.style = "Table Grid"
                    table.autofit = True
                    for ri, r in enumerate(rows):
                        for ci in range(ncols):
                            _fill_cell(table.cell(ri, ci),
                                       r[ci] if ci < len(r) else "",
                                       header=(ri == 0))
                    doc.add_paragraph("")
                continue

            if not s:
                i += 1; continue

            if s.startswith("# "):
                if not first_h1_skipped:   # 首个 H1 与封面重复，跳过
                    first_h1_skipped = True; i += 1; continue
                h = doc.add_heading(s[2:], level=1)
                for r in h.runs: r.font.color.rgb = BRAND
            elif s.startswith("## "):
                h = doc.add_heading(s[3:], level=2)
                for r in h.runs: r.font.color.rgb = BRAND
            elif s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("#### "):
                doc.add_heading(s[5:], level=4)
            elif s.startswith("##### "):
                doc.add_heading(s[6:], level=4)
            elif s == "---" or set(s) == {"─"}:
                pass  # 分隔线略过，靠标题分节
            elif s.startswith("> "):
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(12)
                run = p.add_run(s[2:]); run.italic = True; run.font.color.rgb = GRAY
            elif s.startswith("- ") or s.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                _md_inline_runs(p, s[2:])
            else:
                p = doc.add_paragraph()
                _md_inline_runs(p, s)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return None


# ── HTML 预览导出（P5 可选：贴近 PDF 的卡片/色块观感）────────────────────────
def _html_inline(text: str) -> str:
    import html as _html, re
    t = _html.escape(text)
    t = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", t)
    t = re.sub(r"\*([^*\n]+?)\*", r"<em>\1</em>", t)
    t = t.replace("&lt;br&gt;", "<br>").replace("&lt;br/&gt;", "<br>")
    return t

def export_to_html(product_name: str, content: str) -> str:
    import datetime as _dt
    is_system = "产品体系方案" in content
    subtitle  = "产品体系方案" if is_system else "产品策划方案"

    lines = content.split("\n")
    i, n = 0, len(lines)
    first_h1_skipped = False
    body: list = []

    while i < n:
        s = lines[i].strip()

        # 表格块
        if s.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = [c for c in (_split_row(b) for b in block) if not _is_table_sep(c)]
            rows = [r for r in rows if r]
            if rows:
                ncols = max(len(r) for r in rows)
                t = ['<div class="tbl-wrap"><table>']
                for ri, r in enumerate(rows):
                    tag = "th" if ri == 0 else "td"
                    cells = "".join(
                        f"<{tag}>{_html_inline(r[ci] if ci < len(r) else '')}</{tag}>"
                        for ci in range(ncols)
                    )
                    t.append(f"<tr>{cells}</tr>")
                t.append("</table></div>")
                body.append("".join(t))
            continue

        if not s:
            i += 1; continue

        # 连续 bullet 合成 ul
        if s.startswith("- ") or s.startswith("* "):
            items = []
            while i < n and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                items.append(lines[i].strip()[2:]); i += 1
            body.append("<ul>" + "".join(f"<li>{_html_inline(it)}</li>" for it in items) + "</ul>")
            continue

        if s.startswith("# "):
            if not first_h1_skipped:
                first_h1_skipped = True; i += 1; continue
            body.append(f"<h1>{_html_inline(s[2:])}</h1>")
        elif s.startswith("## "):
            body.append(f"<h2>{_html_inline(s[3:])}</h2>")
        elif s.startswith("### "):
            body.append(f"<h3>{_html_inline(s[4:])}</h3>")
        elif s.startswith("#### "):
            body.append(f"<h4>{_html_inline(s[5:])}</h4>")
        elif s.startswith("##### "):
            body.append(f"<h5>{_html_inline(s[6:])}</h5>")
        elif s == "---" or set(s) == {"─"}:
            pass
        elif s.startswith("> "):
            body.append(f"<blockquote>{_html_inline(s[2:])}</blockquote>")
        else:
            body.append(f"<p>{_html_inline(s)}</p>")
        i += 1

    html_body = "\n".join(body)
    date_str = _dt.date.today().strftime("%Y年%m月")
    css = """
    :root{--brand:#C8102E;--ink:#222;--muted:#888;--line:#eee;--bg:#f5f5f7;}
    *{box-sizing:border-box;}
    body{margin:0;background:var(--bg);color:var(--ink);
      font-family:'Microsoft YaHei','PingFang SC','Helvetica Neue',Arial,sans-serif;line-height:1.75;}
    .page{max-width:860px;margin:24px auto;background:#fff;border-radius:14px;
      box-shadow:0 6px 30px rgba(0,0,0,.08);overflow:hidden;}
    .cover{background:linear-gradient(135deg,var(--brand),#9b0c22);color:#fff;padding:48px 40px;}
    .cover .tag{display:inline-block;border:1px solid rgba(255,255,255,.6);border-radius:20px;
      padding:2px 14px;font-size:13px;margin-bottom:18px;}
    .cover h1{margin:0 0 4px;font-size:32px;font-weight:800;letter-spacing:1px;
      color:#fff;border:none;padding:0;line-height:1.3;text-shadow:0 1px 2px rgba(0,0,0,.18);}
    .cover .sub{margin-top:10px;font-size:16px;opacity:.95;color:#fff;}
    .cover .date{margin-top:6px;font-size:13px;opacity:.85;color:#fff;}
    .content{padding:16px 40px 48px;}
    h1{font-size:23px;color:var(--brand);border-bottom:2px solid var(--brand);padding-bottom:6px;margin:30px 0 14px;}
    h2{font-size:20px;color:#fff;background:var(--brand);display:block;
      padding:9px 16px;border-radius:8px;margin:30px 0 14px;}
    h3{font-size:16px;color:var(--brand);margin:20px 0 8px;padding-left:10px;border-left:4px solid var(--brand);}
    h4{font-size:14.5px;color:#333;margin:14px 0 6px;font-weight:700;}
    h5{font-size:13.5px;color:#555;margin:10px 0 4px;}
    p{margin:8px 0;}
    ul{margin:8px 0 8px 4px;padding-left:22px;}
    li{margin:4px 0;}
    blockquote{margin:12px 0;padding:8px 14px;background:#fafafa;border-left:4px solid var(--brand);
      color:var(--muted);font-style:italic;border-radius:0 6px 6px 0;}
    .tbl-wrap{overflow-x:auto;margin:14px 0;}
    table{border-collapse:collapse;width:100%;font-size:13.5px;}
    th{background:var(--brand);color:#fff;font-weight:700;text-align:left;padding:10px 12px;border:1px solid var(--brand);}
    td{padding:9px 12px;border:1px solid var(--line);vertical-align:top;}
    tr:nth-child(even) td{background:#fbfbfc;}
    strong{color:#111;}
    em{color:var(--muted);font-style:normal;font-size:.92em;}
    @media print{body{background:#fff;}.page{box-shadow:none;margin:0;border-radius:0;}}
    """
    return (
        "<!doctype html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
        "<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">"
        f"<title>{_html_inline(product_name)} · {subtitle}</title>"
        f"<style>{css}</style></head><body><div class=\"page\">"
        f"<div class=\"cover\"><span class=\"tag\">优路教育</span>"
        f"<h1>{_html_inline(product_name)}</h1>"
        f"<div class=\"sub\">{subtitle}</div><div class=\"date\">{date_str}</div></div>"
        f"<div class=\"content\">{html_body}</div></div></body></html>"
    )


# ── Streamlit UI ──────────────────────────────────────────────────────────────
def render_sidebar():
    with st.sidebar:
        st.markdown("## 🎯 AI 产品策划智能体")
        st.caption("一期 · 两段式：定位分析 → 方案生成")
        st.divider()

        st.markdown("**系统状态**")
        st.markdown(f"模型：`{config.QWEN_MODEL}`")
        st.markdown(f"API Key：`...{config.DASHSCOPE_API_KEY[-6:] if config.DASHSCOPE_API_KEY else '未配置'}`")
        bocha_ok = config.bocha_enabled()
        st.markdown(f"博查搜索：{'✅ 已配置' if bocha_ok else '❌ 未配置'}")
        st.markdown(f"知识库：`{config.KB_DIR}`")

        st.divider()
        if st.button("🔄 重建知识库索引", use_container_width=True):
            reset_vectorstore()
            ensure_vectorstore.clear()
            vs, doc_n, chunk_n = build_vectorstore()
            if vs and doc_n:
                st.success(f"✅ 完成：{doc_n} 个案例，{chunk_n} 个片段")
            else:
                st.error("知识库为空，请在 data/ 目录添加 .md 案例文件")

        st.divider()
        st.markdown("**📂 已加载案例**")
        if os.path.exists(config.KB_DIR):
            kb_root = Path(config.KB_DIR)
            files = list(kb_root.glob("**/*.md"))
            for f in sorted(files):
                st.markdown(f"- `{f.relative_to(kb_root)}`")

        st.divider()
        if st.button("🗑️ 重新开始", use_container_width=True):
            for k in ["phase", "form_inputs", "strategy_result", "selected_usp_label"]:
                st.session_state.pop(k, None)
            st.rerun()


def render_form() -> Optional[dict]:
    st.markdown("### 📝 第一步：填写产品信息")
    st.caption("必填项完成后 AI 会自动搜竞品、给出定位方向，再由你选择后生成完整方案")

    col1, col2 = st.columns(2)
    with col1:
        product_name = st.text_input("项目/课程名称 *", placeholder="例：2026年一级建造师精品协议班")
        project_type = st.selectbox("项目类型 *", options=config.PROJECT_TYPES)
        target_audience = st.text_area("目标人群描述 *", height=80,
            placeholder="例：25-40岁，有3年工程经验的在职人员，首考或二战考生")
        learning_motivation = st.multiselect("人群学习动机 *",
            options=config.LEARNING_MOTIVATIONS,
            default=[config.LEARNING_MOTIVATIONS[0]])
    with col2:
        core_advantage = st.text_area("核心竞争优势 *", height=80,
            placeholder="例：押题命中率高、协议保障退费、碎片化学习设计")
        advantage_dims = st.multiselect("重点优势维度",
            options=config.ADVANTAGE_DIMS,
            default=["通过率保障", "师资力量", "服务保障"])
        price_tier = st.selectbox("价格定位", options=config.PRICE_TIERS, index=1)
        extra_info = st.text_area("补充说明（选填）", height=68,
            placeholder="如主打节点、特殊政策背景、需避开的表达等")

    # ── 方案模式 & 证书（P2）────────────────────────────────────────────────
    pt_key = config.project_type_key(project_type)
    mcol1, mcol2 = st.columns([1, 2])
    with mcol1:
        plan_mode = st.selectbox("方案模式", options=config.PLAN_MODES,
            help="产品体系方案适合多证书/多班型联合打包，会按证书分别出招生简章、技能课与就业变现")
    with mcol2:
        cert_options = config.CERTS_BY_TYPE.get(pt_key, [])
        certs = st.multiselect(
            "涉及证书（产品体系模式建议多选）",
            options=cert_options,
            default=cert_options if ("体系" in plan_mode and cert_options) else [],
            help="选中的证书会硬命中对应『证书事实卡』，并在简章中分证书出表。该项目类型暂无预置证书时此处为空，按单产品处理。",
        )

    st.divider()
    # ── 竞品 & 我方现状（可选区块）
    # 用 session_state key 确保 expander 收起时值不丢失
    with st.expander("🔍 竞品信息 & 我方现状（选填，填了定位更准）", expanded=False):
        c1, c2 = st.columns(2)
        with c1:
            competitor_name = st.text_input("主要竞品名称",
                key="f_competitor", placeholder="例：中公教育、环球网校、学天教育")
            our_pass_rate = st.text_input("我方效果数据",
                key="f_pass_rate", placeholder="例：历年综合通过率68%，押题命中率75%")
        with c2:
            sales_objections = st.text_area("销售常见异议", height=80,
                key="f_objections", placeholder="例：太贵了、没时间学、之前报过没过")
            product_weaknesses = st.text_input("已知产品不足",
                key="f_weaknesses", placeholder="例：课程更新不及时、答疑响应慢")
        success_cases = st.text_input("成功案例",
            key="f_cases", placeholder="例：二战考生跟学6个月，一次通过全科")

        st.markdown("---")
        st.caption("📦 以下3项影响简章/师资/课程说明的生成质量，有就填，没有AI会推断并标注「待确认」")
        reference_price = st.text_input("参考售价",
            key="f_price", placeholder="例：标准班1680元，协议班2980元，旗舰班4980元")
        course_structure = st.text_input("课程体系概述",
            key="f_course", placeholder="例：基础课50课时→强化课80课时→冲刺课40课时→押题课20课时")
        teacher_info = st.text_area("师资简介", height=72,
            key="f_teacher", placeholder="例：王某某，一建执业+10年教学经验，擅长管理实务；李某某，造价工程师，主讲计量计价")

    st.divider()
    col_btn, col_tip = st.columns([1, 3])
    with col_btn:
        clicked = st.button("🔍 分析竞品 · 生成定位方向", type="primary", use_container_width=True)
    with col_tip:
        st.caption("AI 将搜索竞品动态 + 赛道大盘，给出 3 个差异化 USP 方向供你选择")

    if clicked:
        missing = []
        if not product_name:        missing.append("项目/课程名称")
        if not target_audience:     missing.append("目标人群描述")
        if not core_advantage:      missing.append("核心竞争优势")
        if not learning_motivation: missing.append("人群学习动机")
        if missing:
            st.error(f"⚠️ 请填写必填项：{'、'.join(missing)}")
            return None
        return {
            "product_name":        product_name,
            "project_type":        project_type,
            "plan_mode":           plan_mode,
            "certs":               certs,
            "target_audience":     target_audience,
            "learning_motivation": "、".join(learning_motivation),
            "core_advantage":      core_advantage,
            "advantage_dims":      "、".join(advantage_dims),
            "price_tier":          price_tier,
            "extra_info":          extra_info or "无",
            "competitor_name":     st.session_state.get("f_competitor", ""),
            "our_pass_rate":       st.session_state.get("f_pass_rate", "暂无数据") or "暂无数据",
            "sales_objections":    st.session_state.get("f_objections", "暂无") or "暂无",
            "product_weaknesses":  st.session_state.get("f_weaknesses", "暂无") or "暂无",
            "success_cases":       st.session_state.get("f_cases", "暂无") or "暂无",
            "reference_price":     st.session_state.get("f_price", "") or "",
            "course_structure":    st.session_state.get("f_course", "") or "",
            "teacher_info":        st.session_state.get("f_teacher", "") or "",
        }
    return None


def render_usp_selection(strategy_result: dict, form_inputs: dict) -> Optional[str]:
    """渲染第二步：展示竞品分析 + USP方向选项，让用户选择"""
    analysis = strategy_result.get("strategy_analysis", "")
    product_name = form_inputs.get("product_name", "")

    st.divider()
    st.markdown(f"### 🧭 第二步：选择定位方向 — {product_name}")
    st.caption("AI 已完成竞品分析，请选择最符合产品实际的 USP 方向，之后生成完整策划方案")

    # 博查搜索诊断 + 原始数据透视
    raw_bocha = strategy_result.get("bocha_results", "")
    bocha_n   = strategy_result.get("bocha_count", 0)
    if bocha_n and bocha_n > 0:
        st.caption(f"🔎 博查搜索返回 {bocha_n} 条结果（已注入竞品分析）")
    else:
        st.warning("⚠️ 博查本次返回 0 条结果——竞品分析已改用内置竞品库兜底。"
                   "若多次为 0，请检查 .env 里的 BOCHA_API_KEY 是否有效/有额度。")
    if raw_bocha and "未获取到竞品搜索数据" not in raw_bocha:
        with st.expander("🔗 查看原始搜索数据（AI分析的原料来源）", expanded=False):
            st.caption("以下是博查搜索返回的原始结果，AI 的竞品洞察即基于此生成")
            st.markdown(raw_bocha)

    # AI 分析报告（展开）
    with st.expander("📊 查看完整竞品分析报告", expanded=True):
        st.markdown(analysis)

    st.divider()

    # 解析出3个方向供选择
    usp_options = _parse_usp_options(analysis)

    if not usp_options:
        st.warning("方向解析失败，请直接在下方输入你希望的定位方向")
        custom = st.text_area("手动输入定位方向", height=100,
            placeholder='例：主打通过率保障，核心主张「押题命中率行业第一，不过退款」')
        col1, _ = st.columns([1, 3])
        with col1:
            if st.button("✅ 用此方向生成方案", type="primary", use_container_width=True):
                return custom
        return None

    # 方向选择
    st.markdown("**请选择一个 USP 定位方向：**")
    option_labels = [f"{opt['label']}  —  {opt['tagline']}" for opt in usp_options]
    option_labels.append("✏️ 自定义方向（手动输入）")

    selected_label = st.radio(
        "USP 方向",
        options=option_labels,
        label_visibility="collapsed",
    )
    st.session_state["selected_usp_label"] = selected_label

    # 自定义
    custom_usp = ""
    if selected_label == option_labels[-1]:
        custom_usp = st.text_area("输入你的定位方向", height=100,
            placeholder='例：主打碎片化学习，核心主张「在职备考不用熬夜，15分钟也能学」')

    st.divider()
    col_btn, col_tip = st.columns([1, 3])
    with col_btn:
        gen_clicked = st.button("🚀 用此方向生成完整方案", type="primary", use_container_width=True)
    with col_tip:
        st.caption("AI 将围绕选定方向生成九章产品策划方案")

    if gen_clicked:
        if selected_label == option_labels[-1]:
            if not custom_usp.strip():
                st.error("请填写自定义方向内容")
                return None
            return custom_usp
        else:
            # 找到对应选项的完整内容
            idx = option_labels.index(selected_label)
            return usp_options[idx]["full_content"]

    return None


def _parse_usp_options(analysis: str) -> list[dict]:
    """从 strategy_analysis 中解析出3个USP方向（多种格式兼容）"""
    import re
    options = []

    # 策略：先按 "方向A/B/C" 分割整段文本，再逐段提取
    # 支持格式：
    #   #### 方向A：【名字】
    #   #### 方向A：名字
    #   **方向A**：名字
    #   方向A：名字
    split_pattern = r"(?=#{1,4}\s*方向\s*[ABC]|(?<!\w)\*\*方向\s*[ABC]\*\*)"
    segments = re.split(split_pattern, analysis)

    for seg in segments:
        # 提取方向字母和名称
        header = re.match(
            r"#{0,4}\s*\*{0,2}方向\s*([ABC])\*{0,2}[：:]\s*【?(.+?)】?\s*\n",
            seg.strip()
        )
        if not header:
            continue
        letter = header.group(1)
        name   = header.group(2).strip().strip("【】")
        content = seg[header.end():]

        # 提取核心主张
        tagline_m = re.search(r"\*{0,2}核心主张\*{0,2}[：:]\s*(.+)", content)
        tagline   = tagline_m.group(1).strip() if tagline_m else ""

        # 提取Slogan预览（备用）
        slogan_m = re.search(r"\*{0,2}Slogan预览\*{0,2}[：:]\s*(.+)", content)
        slogan   = slogan_m.group(1).strip() if slogan_m else ""

        display = tagline[:35] if tagline else (slogan[:35] if slogan else "")
        options.append({
            "label":        f"方向{letter}：{name}",
            "tagline":      display,
            "full_content": f"方向{letter}：{name}\n{content.strip()}",
        })

    # 去重（按字母）
    seen, deduped = set(), []
    for opt in options:
        key = opt["label"][2]  # A/B/C
        if key not in seen:
            seen.add(key)
            deduped.append(opt)

    return deduped


def render_plan_result(plan_result: dict, product_name: str):
    """渲染第三步：完整产品策划方案"""
    plan  = plan_result.get("product_plan", "")
    cases = plan_result.get("retrieved_cases", [])

    st.divider()
    st.markdown(f"## 📄 {product_name} · 产品策划方案")

    c1, c2, c3 = st.columns(3)
    c1.metric("生成状态", "✅ 完成")
    c2.metric("参考案例", f"{len(cases)} 个片段")
    c3.metric("包含章节", "9 个")

    tabs = st.tabs([
        "📋 完整方案",
        "🏷️ Slogan",
        "📖 产品背景",
        "👥 适合人群",
        "💪 产品优势",
        "📑 产品简章",
        "🎓 师资介绍",
        "📚 课程说明",
        "💬 产介话术",
        "🏗️ 搭建思路",
        "📂 参考案例",
    ])

    with tabs[0]:
        st.markdown(plan)
        st.divider()
        dl1, dl2, dl3 = st.columns(3)
        with dl1:
            st.download_button("⬇️ 下载 Markdown", data=plan.encode("utf-8"),
                file_name=f"{product_name}_产品策划方案.md", mime="text/markdown",
                use_container_width=True)
        with dl2:
            word_bytes = export_to_word(product_name, plan)
            if word_bytes:
                st.download_button("⬇️ 下载 Word 文档", data=word_bytes,
                    file_name=f"{product_name}_产品策划方案.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, type="primary")
        with dl3:
            html_str = export_to_html(product_name, plan)
            st.download_button("⬇️ 下载 HTML 预览", data=html_str.encode("utf-8"),
                file_name=f"{product_name}_产品策划方案.html", mime="text/html",
                use_container_width=True)
        with st.expander("👁️ HTML 预览（卡片版观感）", expanded=False):
            components.html(export_to_html(product_name, plan), height=600, scrolling=True)

    sections = _extract_sections(plan)
    with tabs[1]:
        st.markdown("### 主卖点 Slogan 候选")
        st.markdown(sections.get("slogan", "_请查看完整方案_"))
    with tabs[2]:
        st.markdown("### 产品背景")
        st.markdown(sections.get("background", "_请查看完整方案_"))
    with tabs[3]:
        st.markdown("### 适合人群分析")
        st.markdown(sections.get("audience", "_请查看完整方案_"))
    with tabs[4]:
        st.markdown("### 产品优势")
        st.markdown(sections.get("advantage", "_请查看完整方案_"))
    with tabs[5]:
        st.markdown("### 产品简章")
        st.markdown(sections.get("brief", "_请查看完整方案_"))
    with tabs[6]:
        st.markdown("### 师资介绍")
        st.markdown(sections.get("teachers", "_请查看完整方案_"))
    with tabs[7]:
        st.markdown("### 课程说明")
        st.markdown(sections.get("course", "_请查看完整方案_"))
    with tabs[8]:
        st.markdown("### 产介话术")
        st.markdown(sections.get("pitch", "_请查看完整方案_"))
    with tabs[9]:
        st.markdown("### 产品搭建整体思路")
        st.markdown(sections.get("build_logic", "_请查看完整方案_"))
    with tabs[10]:
        if cases:
            for i, doc in enumerate(cases, 1):
                with st.expander(f"参考案例片段 {i}"):
                    st.markdown(doc)
        else:
            st.info("本次未检索到相关历史案例。在 data/ 目录添加真实案例后质量将显著提升。")


def _extract_sections(plan: str) -> dict:
    """
    从 AI 生成的方案文本中提取各章节内容。
    对应 PRODUCT_PLAN_PROMPT 的 9 章节结构。
    """
    import re
    sections = {}
    # 匹配模式：支持"## 主卖点 Slogan"（顶部无编号）以及"## 一、"…"## 八、"带编号格式
    markers = {
        "slogan":      r"##\s*主卖点\s*Slogan",
        "background":  r"##\s*[一1][、.]\s*产品背景",
        "audience":    r"##\s*[二2][、.]\s*适合人群",
        "advantage":   r"##\s*[三3][、.]\s*产品优势",
        "brief":       r"##\s*[四4][、.]\s*产品简章",
        "teachers":    r"##\s*[五5][、.]\s*师资介绍",
        "course":      r"##\s*[六6][、.]\s*课程说明",
        "pitch":       r"##\s*[七7][、.]\s*产介话术",
        "build_logic": r"##\s*[八8][、.]\s*产品搭建",
    }
    keys      = list(markers.keys())
    positions = [m.start() if (m := re.search(p, plan, re.IGNORECASE)) else -1
                 for p in markers.values()]
    for i, (key, pos) in enumerate(zip(keys, positions)):
        if pos == -1:
            continue
        next_pos = next((p for p in positions[i+1:] if p > pos), len(plan))
        sections[key] = plan[pos:next_pos].strip()
    return sections


# ── 主程序 ────────────────────────────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="AI 产品策划智能体",
        page_icon="🎯",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    ok, msg = config.validate_config()
    if not ok:
        st.error(f"❌ {msg}")
        st.code("# 在项目根目录创建 .env：\nDASHSCOPE_API_KEY=sk-xxxx\nBOCHA_API_KEY=sk-xxxx")
        st.stop()

    render_sidebar()
    st.title("🎯 AI 产品策划智能体")
    st.caption("两段式流程：竞品分析 → 选定位方向 → 生成完整策划方案")

    # ── 状态机驱动三个阶段 ──────────────────────────────────────────────────
    phase = st.session_state.get("phase", "form")  # form | strategy | plan

    # ─ 阶段1：填表 ─
    if phase == "form":
        form_inputs = render_form()
        if form_inputs:
            with st.spinner("🔍 正在搜索竞品资料 + 生成定位方向，通常需要 20-40 秒..."):
                result = run_strategy(form_inputs)
            st.session_state["form_inputs"]    = form_inputs
            st.session_state["strategy_result"] = result
            st.session_state["phase"]           = "strategy"
            st.rerun()

    # ─ 阶段2：选方向 ─
    elif phase == "strategy":
        strategy_result = st.session_state.get("strategy_result", {})
        form_inputs     = st.session_state.get("form_inputs", {})

        selected_usp = render_usp_selection(strategy_result, form_inputs)
        if selected_usp:
            plan_inputs = {
                **form_inputs,
                "selected_usp": selected_usp,
                "strategy_analysis": strategy_result.get("strategy_analysis", ""),
                "bocha_results":     strategy_result.get("bocha_results", ""),
            }
            with st.spinner(f"⏳ 正在生成完整策划方案，通常需要 30-60 秒..."):
                plan_result = run_plan(plan_inputs)
            st.session_state["plan_result"]  = plan_result
            st.session_state["phase"]        = "plan"
            st.rerun()

    # ─ 阶段3：展示方案 ─
    elif phase == "plan":
        form_inputs = st.session_state.get("form_inputs", {})
        plan_result = st.session_state.get("plan_result", {})
        product_name = form_inputs.get("product_name", "产品")
        render_plan_result(plan_result, product_name)

        # 允许回到选方向重新生成
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            if st.button("↩️ 换个方向重新生成", use_container_width=True):
                st.session_state["phase"] = "strategy"
                st.session_state.pop("plan_result", None)
                st.rerun()
        with col2:
            if st.button("🆕 新建策划方案", use_container_width=True):
                for k in ["phase", "form_inputs", "strategy_result", "plan_result", "selected_usp_label"]:
                    st.session_state.pop(k, None)
                st.rerun()


if __name__ == "__main__":
    main()
